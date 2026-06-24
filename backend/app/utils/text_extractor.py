"""
utils/text_extractor.py
-------------------------
Extracts raw text from uploaded resume files (PDF or DOCX).

Each format has its own extractor so a corrupt or unreadable file
raises a clear, catchable HTTPException instead of crashing the
request with a raw library exception.
"""

from fastapi import HTTPException, status

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover - guarded so the app still imports
    PdfReader = None

try:
    import docx  # python-docx
except ImportError:  # pragma: no cover
    docx = None


def _extract_pdf_text(file_path: str) -> str:
    if PdfReader is None:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="PDF support is not installed (pypdf missing)",
        )
    try:
        reader = PdfReader(file_path)
        pages_text = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages_text).strip()
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Could not read PDF file: {e}",
        )


def _extract_docx_text(file_path: str) -> str:
    if docx is None:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="DOCX support is not installed (python-docx missing)",
        )
    try:
        document = docx.Document(file_path)
        paragraphs = [p.text for p in document.paragraphs]

        # Resumes often use tables for layout — pull that text in too
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        paragraphs.append(cell.text)

        return "\n".join(p for p in paragraphs if p.strip()).strip()
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Could not read DOCX file: {e}",
        )


def extract_text(file_path: str, ext: str) -> str:
    """
    Dispatches to the right extractor based on file extension.
    *ext* should be ".pdf" or ".docx" (lowercase, with the dot).

    Raises HTTPException(422) if no text could be extracted —
    e.g. the resume is a scanned image with no embedded text layer.
    """
    if ext == ".pdf":
        text = _extract_pdf_text(file_path)
    elif ext == ".docx":
        text = _extract_docx_text(file_path)
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"No text extractor available for '{ext}'",
        )

    if not text:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="No extractable text found in the resume (it may be a scanned image)",
        )

    return text
